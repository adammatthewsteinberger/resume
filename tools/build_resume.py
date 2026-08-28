#!/usr/bin/env python3
"""Single-source résumé builder: emits README.md, .txt, .docx, .html (→ PDF via Chrome).

Content is data below; every output is derived from it so the formats never drift.
Sources of truth: LinkedIn profile export (2026-08-18), vibewithadam.matthewsteinberger.com,
GitHub/PyPI, and the book manuscripts. No invented metrics.
"""
import html, json, re, sys, os

OUT = sys.argv[1] if len(sys.argv) > 1 else "."

# ---------------------------------------------------------------- content
NAME = "Adam Matthew Steinberger"
TITLE = "Staff Software Architect & AI Automation Engineer"
TAGLINE = ("RAG systems, event-driven Azure microservices, and automation pipelines "
           "that the people who inherit them can actually run.")
CONTACT = [
    ("Greenville, SC (remote / US remote)", None),
    ("+1-864-517-4117", "tel:+18645174117"),
    ("adam@matthewsteinberger.com", "mailto:adam@matthewsteinberger.com"),
    ("linkedin.com/in/adammatthewsteinberger", "https://www.linkedin.com/in/adammatthewsteinberger/"),
    ("github.com/adammatthewsteinberger", "https://github.com/adammatthewsteinberger"),
    ("vibewithadam.matthewsteinberger.com", "https://vibewithadam.matthewsteinberger.com"),
    ("vibewithadam.matthewsteinberger.com/join-me", "https://vibewithadam.matthewsteinberger.com/join-me"),
]
AVAILABILITY = ("Available from September 2026 for Staff Software Architect, AI Automation Engineer, "
                "Staff/Principal AI Engineer, Solutions Architect, or Platform/Automation Engineer roles "
                "— W2 preferred, contract-to-hire OK.")

SUMMARY = (
    "I build AI systems that actually work inside enterprise environments — production-grade platforms that "
    "handle real data, real security requirements, and real organizational complexity, *not demos*. **13+ years** "
    "across AI, cloud, fintech, insurance, healthcare, and cybersecurity. Most recently: **sole architect of a "
    "multi-vendor AI governance gateway**, **co-lead of a 20-microservice AI payroll platform**, author of two "
    "identity-governance-as-code control planes and the shared Python platform library adopted by 17+ repositories "
    "— all on private AKS with secretless (OIDC / workload identity) delivery and supply-chain security in CI. "
    "The work starts with a business problem nobody has solved cleanly and ends with **something the people who "
    "inherit it can run** — architecture documented before code, junior developers trained in parallel, handoffs "
    "that hold. I document everything for the same reason a RAG pipeline cites its sources."
)

SKILLS = [
    ("AI & LLM Systems", "RAG and vector search (pgvector, FAISS, Azure AI Search, Pinecone), multi-vendor LLM gateways with cost/policy "
                         "governance, AI agents and multi-agent orchestration, MCP servers, structured outputs, HITL gating, LoRA fine-tuning; "
                         "Azure OpenAI/Foundry, Claude, GPT, Gemini, Grok, Mistral, vLLM, Ollama, LangChain, Document Intelligence, "
                         "Content Safety; OWASP LLM Top 10 / NIST AI RMF"),
    ("Azure & Cloud", "AKS (private clusters, workload identity, KEDA), Functions, App Service, Service Bus, Event Hubs, Key Vault, "
                      "App Configuration, App Insights/OpenTelemetry, Data Explorer (KQL), Cosmos DB, PostgreSQL, Redis, Blob/Data Lake, "
                      "App Gateway + WAF, Private Endpoints, Entra ID, Microsoft Graph; sovereign/government cloud; AWS (Textract, Lambda, S3)"),
    ("Platform, DevSecOps & Delivery", "Terraform, Bicep, Helm, Kustomize, Flux/Argo CD GitOps, Docker, GitHub Actions "
                          "(OIDC federated credentials, self-hosted runners), build-once/promote-by-digest; Trivy, Semgrep, "
                          "CodeQL, Bandit, Gitleaks/TruffleHog, Checkov, pip-audit, SBOM (Syft/CycloneDX), Cosign keyless "
                          "signing, Kyverno/OPA admission; threat modeling (STRIDE), SOC 2 readiness, ADRs"),
    ("Identity & Governance", "Okta (core, IGA, Workflows), Microsoft Entra ID, SAML 2.0, OIDC/OAuth 2.0, workload identity "
                              "federation, RBAC (control and data plane), governance-as-code reconciliation, SOX-aligned access "
                              "governance, GxP-classified functional specifications"),
    ("Languages & Frameworks", "Python 3.11/3.12 (FastAPI, Flask, SQLAlchemy 2, Pydantic, kopf), TypeScript/NestJS, "
                               "Next.js 15/16 + React 19, C#/.NET (Web API, MVC), Java Spring Boot, gRPC/REST, OpenAPI 3.1, Bash, SQL, KQL"),
    ("Quality & Data", "pytest, Hypothesis, mutmut, contract/e2e/chaos tests, mypy --strict, ruff, import-linter-enforced onion "
                       "architecture; PostgreSQL, MongoDB, Snowflake, SQL Server, Oracle, Redis; ETL/API integrations (HubSpot, "
                       "SharePoint, Salesforce, Outlook)"),
    ("Delivery & Leadership", "Discovery → documented solution → Jira decomposition → mentored handoff; Scrum (CSM), "
                              "Security-First Scrum (author), evidence-based delivery (DORA/CHAOS/QSM), architecture documentation, mentoring"),
]

EXPERIENCE = [
    {
        "org": "The Vizius Group", "loc": "Greenville, SC",
        "role": "Senior Azure and AI Development Engineer", "dates": "Sep 2025 – Aug 2026",
        "blurb": "Cybersecurity firm. Designed and shipped six production platforms on private AKS for mid-market and enterprise "
                 "clients (manufacturing, health-tech, financial services, a SOX-regulated enterprise), plus the firm's shared "
                 "platform library, architecture practice, and thought-leadership program. ~1,050 commits across 16 repos; sole "
                 "or lead author on 9. Client identities withheld.",
        "bullets": [
            "**AI governance gateway** _(sole architect, ~54k LOC)_ — one policy-enforced, OpenAI-compatible API in front of the "
            "full Azure AI surface plus Anthropic, OpenAI/Codex, Cursor, Grok, and Gemini: per-project allowlists, "
            "allow/deny/fallback policy with hot reload, multi-unit Redis rate limiting, per-call USD cost attribution with "
            "enforced spend caps (denial-of-wallet control), and an HMAC-signed, hash-chained, write-once audit trail; Entra ID "
            "app-role auth over workload identity — **no API keys in the path**. Agent sandboxing, egress policy, and SSRF checks "
            "mapped to OWASP LLM Top 10 / NIST AI RMF; Python SDK, CLI, MCP server, Next.js admin portal; 9 autoscaled pods on "
            "private AKS via GitOps. **Migrated three product teams onto it** and retired their app-held credentials.",
            "**AI payroll automation platform** _(co-lead, ~420k LOC)_ — 20 microservices in an onion-architecture monorepo across "
            "four human-approved phases with the final submission modeled as irreversible; RAG over the document store, "
            "AI-directed spreadsheet corrections, earnings and report review. Owned Terraform, 20 Helm charts, Kustomize, "
            "GitOps, and 10 CI/CD workflows; 585 test modules across unit/integration/contract/e2e/smoke. Architecture "
            "**production-ready at day 45**; a junior developer trained in parallel now owns it.",
            "**Technical report generation platform** _(lead, ~54k LOC)_ — turns raw electrical-testing instrument data into "
            "standards-aware customer deliverables: mail-webhook ingestion with per-document fan-out, a multi-vendor parser seam, "
            "a deterministic deficiency analyzer fed by a scraped standards store plus LLM review, blocking data-quality validation, "
            "SAML 2.0 + Entra dual-issuer SSO, per-user bearer auth replacing a shared API key. **Eliminated silent false-success "
            "deploys**; authored the SOC 2 readiness assessment, threat model, ADRs, and an evidence-based delivery operating model.",
            "**Identity governance as code** _(sole author, two control planes)_ — a Kubernetes operator (kopf) that reconciles "
            "directory governance state against Git-declared custom resources with **fully secretless multi-tenant auth** "
            "(federated credentials, zero stored tenant secrets) and an LLM that drafts pull requests for judgment calls; and "
            "an IdP governance platform managing 40 resource kinds through six addressing patterns with drift "
            "classification (auto-remediate safe, PR + human approval for destructive), point-in-time reversion, and dual "
            "APM/SIEM log shipping. Plus a versioned, idempotent sync API for 114+ directory groups that replaced a low-code workflow.",
            "**Multi-system ticket relay** _(sole author, ~20k LOC)_ — N-way sync with a symmetric schema (no privileged hub): "
            "version vectors, echo suppression, a conflict policy engine that downgrades unimplemented strategies to manual hold, "
            "edge HMAC verification with vault-backed per-tenant secrets, config-driven generic connector. **653 tests, 93% coverage**, "
            "mypy --strict clean, import-linter-enforced pure domain, property/mutation/chaos tests proving convergence.",
            "**Multi-tenant observability portal** _(lead)_ — sub-second streaming, analytical, and federated APM/vendor/cost planes, "
            "**every payload tagged with its freshness**; CLI, REST, HMAC webhooks, MCP server, and SDK over one core; SAML SSO; KEDA.",
            "**vibey-bootstrap** _(formerly azure-bootstrap; open source, MIT)_ — authored the firm's shared Python platform library "
            "through three major versions on PyPI, **adopted by 17+ repositories**: four-phase logging↔config bootstrap, structured "
            "logging with correlation IDs and masking, tiered alerting, ingress classifier, dead-letter-aware consumers, ten logging "
            "transports behind a never-block/never-raise shipper, transactional outbox; 86% coverage; four platforms refactored to "
            "delete the code it replaced.",
            "**DevSecOps, secretless by default** — OIDC workload identity federation across 20 CI workflows in 9 repositories, "
            "managed identity at runtime, CSI-driver vault secrets; supply-chain pipelines with SAST, SCA, IaC scanning, secret "
            "detection, SBOM generation, keyless image signing, and policy-as-code admission; cleared 24 IaC policy findings; "
            "**security self-reviews** closed an auth bypass, path traversal, SSRF, timing-unsafe comparison, query injection, and "
            "an over-scoped CI credential before release. Cross-tenant production migration on OIDC and least-privilege RBAC.",
            "**Architecture & advisory** — **five formal architecture document sets (~180 pages)** including a three-tier package "
            "(43-page design, 10-page executive summary, one-sheet) and a STRIDE threat model; identity-governance advisory for a "
            "**SOX-regulated enterprise of ~5,700 identities** (market survey, platform decision report, API/SDK/MCP coverage assessment "
            "across eight platforms, GxP-classified functional specifications, SOX-to-IAM risk mapping); AI vendor terms comparison "
            "for legal and procurement.",
            "**Enablement & thought leadership** — authored *Security-First Scrum* (framework, two training manuals, four AI-agent "
            "rulesets), an evidence-based delivery velocity playbook, and a **~110,000-word technical reference library** compiled "
            "into vibey-skills (18 plugins / 71 skills); mentored junior developers on three projects; built the firm's LinkedIn "
            "thought-leadership program end to end, including a narrative white paper on export-control compliance and cloud "
            "enclave architecture produced and written from recorded expert interviews.",
        ],
    },
    {
        "org": "The Apologist Project (volunteer)", "loc": "Remote",
        "role": "Volunteer Software Architect — open-source-style contribution", "dates": "Apr 2026 – Present",
        "blurb": "Nonprofit AI apologetics chat platform (Next.js seeker app + Laravel/Filament backend). Unpaid, concurrent with Vizius.",
        "bullets": [
            "**Project Excite** — designed and built an **adapter-based relay microservice** handing seekers from the AI to live volunteers "
            "on Chatwoot or EchoGlobal: abstract adapter + concrete adapters, **explicit session state machine with idempotent teardown**, "
            "Redis-backed session manager, HMAC-verified webhooks, QStash-queued delivery, shared in-session @agent. Three technical "
            "executive summaries and a unified relay schema reference written before implementation.",
            "**Shipped across split PR stacks** (schema, relay lib/HTTP, backend proxy, client UI, admin monitoring) plus **security hardening** "
            "(XSS via DOMPurify, CORS allowlist, Sentry PII off, rate limiting) and CI/PHPUnit repair; ~68 commits.",
        ],
    },
    {
        "org": "Adam Matthew Steinberger LLC", "loc": "Greenville, SC",
        "role": "Independent product design", "dates": "2026",
        "bullets": [
            "**Business plans and software architecture documents** for two SaaS concepts — a mobile-first social platform (React Native, "
            "FastAPI, Azure Container Apps) and a decentralized confidential-AI protocol.",
        ],
    },
    {
        "org": "Adam Matthew Steinberger LLC", "loc": "Greenville, SC",
        "role": "Senior Software Engineering Consultant", "dates": "Mar 2025 – Aug 2025",
        "blurb": "Four engagements in six months.",
        "bullets": [
            "**Self-hosted RAG chatbot** _(non-profit)_ — on-premise RAG on Mistral-7B + FAISS behind an OpenAI-compatible vLLM API; "
            "Grafana/Prometheus on every token; Docker on bare metal; **zero external dependencies; shipped in 30 days**.",
            "**Cloud RAG chatbot** _(sales agency)_ — Gemini-based RAG with API-driven web search; **shipped in 30 days**.",
            "**Web push notification system** _(non-profit, GodFocus)_ — timezone-aware scheduling, personalization, VAPID encryption; "
            "**159/159 tests, 85.84% coverage** via AI-assisted TDD in **5 billable hours against a 30+ hour estimate**.",
            "**Codebase review & architecture** _(non-profit)_ — **190+ files / 59,000 lines in 10 hours**; surfaced 5% test coverage and "
            "missing auth middleware; delivered a technical brief, executive summary, and phased Onion roadmap.",
        ],
    },
    {
        "org": "Lima One Capital", "loc": "Greenville, SC",
        "role": "Senior Software Engineer", "dates": "May 2023 – Feb 2025",
        "blurb": "Real-estate lender.",
        "bullets": [
            "**Rearchitected the core integration layer** from legacy Mulesoft APIs into NestJS microservices (gRPC + REST) on PostgreSQL.",
            "**Full-stack .NET/React** work on a mortgage-broker platform: credit-report integrations and pricing-engine APIs.",
            "**ETL pipelines and API connectors** across HubSpot, SharePoint, Snowflake, Salesforce, and third-party providers.",
            "**Built Snow Portal**, a Snowflake job scheduler that **replaced Alteryx at a fraction of the cost**; automated HR-to-ITSM sync.",
        ],
    },
    {
        "org": "Transcat", "compact": True, "loc": "Rochester, NY",
        "role": "Senior Software Engineer", "dates": "Apr 2022 – Jan 2023",
        "bullets": [
            "*Led a team* delivering .NET Web APIs and a React front end for lab-equipment calibration; hardened the Magento channel.",
        ],
    },
    {
        "org": "LeaseTrack", "compact": True, "loc": "Latham, NY",
        "role": "Senior Software Engineer", "dates": "Jun 2021 – Apr 2022",
        "bullets": [
            "Python + AWS Textract insurance-document parsing, plus a Java Spring Boot annotation system feeding the ML training pipeline.",
        ],
    },
    {
        "org": "Akmazio Software", "compact": True, "loc": "Albany, NY",
        "role": "Senior Software Engineer (founding engineer)", "dates": "May 2020 – May 2021",
        "bullets": [
            "*Founding engineer:* built the entire C#/.NET + MS SQL backend (DigitalOcean) for an advisor-matching platform; wrote the business plan, managed interns and a 1099 developer, ran a distributed Scrum test team.",
        ],
    },
    {
        "org": "Bestpass by Fleetworthy", "compact": True, "loc": "Albany, NY",
        "role": "Software Engineer", "dates": "Sep 2019 – Apr 2020",
        "bullets": [
            "Toll-billing system in C# MVC + Knockout.js; *introduced automated unit testing* to a legacy codebase that had none.",
        ],
    },
    {
        "org": "New York State Insurance Fund (NYSIF)", "compact": True, "loc": "Albany, NY",
        "role": "Software Engineer", "dates": "Mar 2015 – Aug 2019",
        "bullets": [
            "*Migrated VB6 systems to C# MVC*, refactored Oracle EDI integrations, mentored juniors, standardized engineering process.",
        ],
    },
    {
        "org": "Town and Country Computer Services", "compact": True, "loc": "Schenectady, NY",
        "role": "Junior Software Engineer", "dates": "Jul 2013 – Mar 2015",
        "bullets": [
            "C# ASP.NET / SQL Server quoting, rating, and reporting apps used all day by insurance underwriters; *client-facing from day one*.",
        ],
    },
    {
        "org": "GE HealthCare", "compact": True, "loc": "Barrington, IL",
        "role": "Junior Software Engineer", "dates": "Aug 2012 – Feb 2013",
        "bullets": [
            "Zero Footprint (ZFP), a browser-based JavaScript CT/MRI viewer for real-time 3D scrolling; *built the full i18n feature*. First Scrum team.",
        ],
    },
]

OPEN_SOURCE = [
    ("claudeloop · codexloop · cursorloop · agyloop · qwenloop · vibey · vibey-gh",
     "Onion-architected autonomous session runners for Claude Code, OpenAI Codex, Cursor Agent, Google Antigravity, and a fully local "
     "Qwen 2.5 Coder _(never block on a human; distinguish rate-limit windows from exhausted credits)_; **vibey**, the six-phase "
     "PostgreSQL-backed conductor on top of them; and **vibey-gh**, stdlib-only release automation (provenance, merge train, "
     "dual-channel releases).", "https://github.com/adammatthewsteinberger/claudeloop"),
    ("vibey-bootstrap · vibey-skills",
     "The Azure Functions cross-cutting layer _(formerly azure-bootstrap; 17+ repos)_ and a Claude Code marketplace of evidence-grounded "
     "practitioner skills _(formerly vibe-engineering-skills)_.", "https://github.com/adammatthewsteinberger/vibey-bootstrap"),
]
# Every " · "-separated package name in an OPEN_SOURCE entry is linked to its own repo.
OSS_LINKS = {n: f"https://github.com/adammatthewsteinberger/{n}" for n in
             ["claudeloop", "codexloop", "cursorloop", "agyloop", "qwenloop", "vibey", "vibey-gh",
              "vibey-bootstrap", "vibey-skills"]}
OSS_NOTE = "All MIT-licensed, on PyPI. Contributors and volunteers welcome — "
OSS_NOTE_URL = "https://vibewithadam.matthewsteinberger.com/join-me"
OSS_NOTE_LINK_TEXT = "vibewithadam.matthewsteinberger.com/join-me"

def oss_names(name):
    """Split 'a · b · c' into [(name, url-or-None), ...]."""
    return [(n, OSS_LINKS.get(n)) for n in name.split(" · ")]

PUBLICATIONS = [
    ("Novice to Navigator: Your Guide to AI Chatbots for Business",
     "Plain-English guide to RAG chatbots for decision-makers; **first edition free online**, second edition in development _(ISBN 979-8274310628)_.", "https://vibewithadam.matthewsteinberger.com/novice-to-navigator"),
]

EDUCATION = [
    ("Skidmore College", "B.A., Computer Science", "2010 – 2012"),
    ("Rensselaer Polytechnic Institute", "Electrical and Electronics Engineering", "2008 – 2010"),
]
CERTS = [("Certified ScrumMaster (CSM)", "Scrum Alliance", "2021")]

# ---------------------------------------------------------------- inline emphasis
# Content strings may carry lightweight inline markup:
#   **bold**  – scannable handles and the one or two outcomes worth catching at a glance
#   *italic*  – ordinary emphasis (titles, a key verb in a one-liner)
#   _muted_   – de-emphasized qualifiers such as (sole architect, ~54k LOC); gray italic
# Markdown passes it through (`_x_` is italic there too); TXT strips it; HTML and DOCX render real runs.
_EM = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_")

def plain(s):
    """Strip inline markup → plain text."""
    return _EM.sub(lambda m: next(g for g in m.groups() if g is not None), s)

def spans(s):
    """Yield (text, style) runs for a marked-up string; style ∈ {None, 'b', 'i', 'muted'}."""
    pos = 0
    for m in _EM.finditer(s):
        if m.start() > pos:
            yield s[pos:m.start()], None
        if m.group(1) is not None:   yield m.group(1), "b"
        elif m.group(2) is not None: yield m.group(2), "i"
        else:                        yield m.group(3), "muted"
        pos = m.end()
    if pos < len(s):
        yield s[pos:], None

def rich(s):
    """Marked-up string → escaped HTML with <b>/<i>/<i class=muted>."""
    out = []
    for text, style in spans(s):
        t = html.escape(text, quote=False)
        if style == "b":       t = f"<b>{t}</b>"
        elif style == "i":     t = f"<i>{t}</i>"
        elif style == "muted": t = f"<i class='muted'>{t}</i>"
        out.append(t)
    return "".join(out)

# ---------------------------------------------------------------- markdown
def md():
    L = []
    L.append(f"# {NAME}")
    L.append("")
    L.append(f"**{TITLE}** — {TAGLINE}")
    L.append("")
    parts = []
    for text, href in CONTACT:
        parts.append(f"[{text}]({href})" if href else text)
    L.append(" · ".join(parts))
    L.append("")
    L.append(f"> {AVAILABILITY}")
    L.append("")
    L.append("## Summary")
    L.append("")
    L.append(SUMMARY)
    L.append("")
    L.append("## Core Skills")
    L.append("")
    for k, v in SKILLS:
        L.append(f"- **{k}:** {v}")
    L.append("")
    L.append("## Experience")
    L.append("")
    for e in [x for x in EXPERIENCE if not x.get("compact")]:
        L.append(f"### {e['org']} — {e['role']}")
        L.append(f"*{e['dates']} · {e['loc']}*")
        L.append("")
        for b in e["bullets"]:
            L.append(f"- {b}")
        L.append("")
    L.append("### Earlier experience")
    L.append("")
    for e in [x for x in EXPERIENCE if x.get("compact")]:
        L.append(f"- **{e['org']}** — {e['role']}, {e['loc']} ({e['dates']}). {' '.join(e['bullets'])}")
    L.append("")
    L.append("## Open Source")
    L.append("")
    for name, desc, url in OPEN_SOURCE:
        linked = " · ".join(f"[{n}]({u})" if u else n for n, u in oss_names(name))
        L.append(f"- **{linked}** — {desc}")
    L.append("")
    L.append(f"{OSS_NOTE}[{OSS_NOTE_LINK_TEXT}]({OSS_NOTE_URL})")
    L.append("")
    L.append("## Publications")
    L.append("")
    for name, desc, url in PUBLICATIONS:
        L.append(f"- **[{name}]({url})** — {desc}")
    L.append("")
    L.append("## Education & Certifications")
    L.append("")
    for school, deg, yrs in EDUCATION:
        L.append(f"- **{school}** — {deg} ({yrs})")
    for c, org, yr in CERTS:
        L.append(f"- **{c}** — {org} ({yr})")
    L.append("")
    L.append("---")
    L.append("")
    L.append("Formats: [PDF](adam-steinberger-resume.pdf) · [DOCX](adam-steinberger-resume.docx) · "
             "[TXT](adam-steinberger-resume.txt) · [Scrum certificate](scrum-certificate.pdf) · "
             "Everything else: [vibewithadam.matthewsteinberger.com/hire-me](https://vibewithadam.matthewsteinberger.com/hire-me)")
    L.append("")
    L.append("License: code [MIT](LICENSE) · résumé content [CC BY 4.0](LICENSE-CONTENT.md) · builder: `tools/build_resume.py`")
    L.append("")
    return "\n".join(L)

# ---------------------------------------------------------------- plain text
def txt():
    L = [NAME.upper(), TITLE, TAGLINE, ""]
    L.append(" | ".join(t for t, _ in CONTACT))
    L.append("")
    L.append(AVAILABILITY)
    L.append("")
    L.append("SUMMARY"); L.append(SUMMARY); L.append("")
    L.append("CORE SKILLS")
    for k, v in SKILLS:
        L.append(f"- {k}: {v}")
    L.append("")
    L.append("EXPERIENCE")
    for e in [x for x in EXPERIENCE if not x.get("compact")]:
        L.append("")
        L.append(f"{e['org']} | {e['loc']}")
        L.append(f"{e['role']} | {e['dates']}")
        for b in e["bullets"]:
            L.append(f"- {b}")
    L.append("")
    L.append("EARLIER EXPERIENCE")
    for e in [x for x in EXPERIENCE if x.get("compact")]:
        L.append(f"- {e['org']} - {e['role']}, {e['loc']} ({e['dates']}). {' '.join(e['bullets'])}")
    L.append("")
    L.append("OPEN SOURCE")
    for name, desc, url in OPEN_SOURCE:
        L.append(f"- {name}: {desc} ({url})")
    L.append(OSS_NOTE + OSS_NOTE_LINK_TEXT)
    L.append("")
    L.append("PUBLICATIONS")
    for name, desc, url in PUBLICATIONS:
        L.append(f"- {name}: {desc} ({url})")
    L.append("")
    L.append("EDUCATION & CERTIFICATIONS")
    for school, deg, yrs in EDUCATION:
        L.append(f"- {school} - {deg} ({yrs})")
    for c, org, yr in CERTS:
        L.append(f"- {c} - {org} ({yr})")
    L.append("")
    return plain("\n".join(L))

# ---------------------------------------------------------------- html (→ pdf)
def h(s): return html.escape(s, quote=False)

def html_doc():
    css = """
    @page { size: Letter; margin: 0.32in 0.42in; }
    * { box-sizing: border-box; }
    body { font-family: -apple-system, "Helvetica Neue", Helvetica, Arial, sans-serif; color: #1f2328; font-size: 8.5pt; line-height: 1.12; margin: 0; }
    h1 { font-size: 19pt; margin: 0 0 1pt; letter-spacing: -0.3px; }
    .title { font-size: 12pt; font-weight: 600; color: #0969da; margin: 0 0 3pt; }
    .tagline { color: #57606a; margin: 0 0 6pt; }
    .contact { font-size: 9.2pt; color: #24292f; margin: 0 0 6pt; }
    .contact a { color: #0969da; text-decoration: none; }
    .avail { font-size: 9pt; background: #dafbe1; border: 1px solid #1a7f37; color: #116329; border-radius: 6px; padding: 5px 8px; margin: 0 0 8pt; }
    h2 { font-size: 10.5pt; text-transform: uppercase; letter-spacing: 0.6px; color: #0969da; border-bottom: 1.5px solid #d0d7de; padding-bottom: 2pt; margin: 5pt 0 2.5pt; }
    p { margin: 0 0 4pt; }
    ul { margin: 1.5pt 0 3pt; padding-left: 13pt; }
    li { margin: 0 0 0.8pt; }
    .job { margin: 0 0 3pt; }
    .job-head, .role, .blurb { break-after: avoid; page-break-after: avoid; }
    h2 { break-after: avoid; page-break-after: avoid; }
    .job-head { display: flex; justify-content: space-between; align-items: baseline; }
    .job-head .org { font-weight: 700; font-size: 10.6pt; }
    .job-head .dates { color: #57606a; font-size: 9.4pt; white-space: nowrap; }
    .role { color: #24292f; font-style: italic; margin: 0 0 2pt; }
    .blurb { color: #57606a; margin: 0 0 2pt; }
    .skills li { margin-bottom: 1.6pt; }
    .early { columns: 1; }
    a { color: #0969da; text-decoration: none; }
    li b, p b { font-weight: 700; color: #0b1320; }
    i.muted { color: #57606a; }
    """
    L = [f"<!doctype html><html><head><meta charset='utf-8'><title>{h(NAME)} — Résumé</title><style>{css}</style></head><body>"]
    L.append(f"<h1>{h(NAME)}</h1><div class='title'>{h(TITLE)}</div><div class='tagline'>{h(TAGLINE)}</div>")
    L.append("<div class='contact'>" + " &nbsp;·&nbsp; ".join(
        (f"<a href='{href}'>{h(t)}</a>" if href else h(t)) for t, href in CONTACT) + "</div>")
    L.append(f"<div class='avail'>{h(AVAILABILITY)}</div>")
    L.append(f"<h2>Summary</h2><p class='summary'>{rich(SUMMARY)}</p>")
    L.append("<h2>Core Skills</h2><ul class='skills'>" + "".join(f"<li><b>{h(k)}:</b> {rich(v)}</li>" for k, v in SKILLS) + "</ul>")
    L.append("<h2>Experience</h2>")
    for e in [x for x in EXPERIENCE if not x.get("compact")]:
        L.append("<div class='job'>")
        L.append(f"<div class='job-head'><span class='org'>{h(e['org'])} <span style='font-weight:400;color:#57606a'>· {h(e['loc'])}</span></span><span class='dates'>{h(e['dates'])}</span></div>")
        L.append(f"<div class='role'>{h(e['role'])}</div>")
        L.append("<ul>" + "".join(f"<li>{rich(b)}</li>" for b in e["bullets"]) + "</ul></div>")
    L.append("<div class='role' style='font-style:normal;font-weight:700;margin-top:4pt'>Earlier experience</div><ul>")
    for e in [x for x in EXPERIENCE if x.get("compact")]:
        L.append(f"<li><b>{h(e['org'])}</b> — {h(e['role'])}, {h(e['loc'])} <span style='color:#57606a'>({h(e['dates'])})</span>. {rich(' '.join(e['bullets']))}</li>")
    L.append("</ul>")
    L.append("<h2>Open Source &amp; Publications</h2><ul>" + "".join("<li><b>" + " · ".join(f"<a href='{lu}'>{h(ln)}</a>" if lu else h(ln) for ln, lu in oss_names(n)) + f"</b> — {rich(d)}</li>" for n, d, u in OPEN_SOURCE) + f"<li>{h(OSS_NOTE)}<a href='{OSS_NOTE_URL}'>{h(OSS_NOTE_LINK_TEXT)}</a></li>" + "".join(f"<li><b><a href='{u}'>{h(n)}</a></b> — {rich(d)}</li>" for n, d, u in PUBLICATIONS) + "</ul>")
    L.append("<h2>Education &amp; Certifications</h2><p>" +
             " &nbsp;·&nbsp; ".join([f"<b>{h(s)}</b> — {h(d)} ({h(y)})" for s, d, y in EDUCATION] +
                                   [f"<b>{h(c)}</b> — {h(o)} ({h(y)})" for c, o, y in CERTS]) + "</p>")
    L.append("</body></html>")
    return "".join(L)

# ---------------------------------------------------------------- docx
def docx_doc(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.65)
    base = d.styles["Normal"]; base.font.name = "Calibri"; base.font.size = Pt(10.5)
    base.paragraph_format.space_after = Pt(2)

    def add_hyperlink(par, url, text):
        part = par.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color"); c.set(qn("w:val"), "0969DA"); rPr.append(c)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        new_run.append(rPr); t = OxmlElement("w:t"); t.text = text; new_run.append(t); hl.append(new_run)
        par._p.append(hl)

    def heading(text):
        p = d.add_paragraph(); r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0x09, 0x69, 0xDA); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        # bottom border
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:space"), "1"); b.set(qn("w:color"), "D0D7DE")
        pbdr.append(b); pPr.append(pbdr)

    def add_rich(par, text):
        for t, style in spans(text):
            r = par.add_run(t)
            if style == "b": r.bold = True
            elif style == "i": r.italic = True
            elif style == "muted": r.italic = True; r.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)

    def bullet(text, bold_prefix=None, url=None):
        p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1.5)
        if bold_prefix:
            if url:
                add_hyperlink(p, url, bold_prefix); p.add_run(" — "); add_rich(p, text)
            else:
                r = p.add_run(bold_prefix + ": "); r.bold = True; add_rich(p, text)
        else:
            add_rich(p, text)

    p = d.add_paragraph(); r = p.add_run(NAME); r.bold = True; r.font.size = Pt(20)
    p = d.add_paragraph(); r = p.add_run(TITLE); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x09, 0x69, 0xDA)
    p = d.add_paragraph(); r = p.add_run(TAGLINE); r.italic = True; r.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)
    p = d.add_paragraph()
    for i, (t, href) in enumerate(CONTACT):
        if i: p.add_run("  ·  ")
        if href: add_hyperlink(p, href, t)
        else: p.add_run(t)
    p = d.add_paragraph(); r = p.add_run(AVAILABILITY); r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x11, 0x63, 0x29)

    heading("Summary"); add_rich(d.add_paragraph(), SUMMARY)
    heading("Core Skills")
    for k, v in SKILLS: bullet(v, bold_prefix=k)
    heading("Experience")
    for e in [x for x in EXPERIENCE if not x.get("compact")]:
        p = d.add_paragraph(); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{e['org']}"); r.bold = True; r.font.size = Pt(11)
        p.add_run(f"  ·  {e['loc']}")
        p.add_run(f"\t{e['dates']}").font.color.rgb = RGBColor(0x57, 0x60, 0x6A)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), alignment=2)
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(1); r = p.add_run(e["role"]); r.italic = True
        for b in e["bullets"]: bullet(b)
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(5); r = p.add_run("Earlier experience"); r.bold = True
    for e in [x for x in EXPERIENCE if x.get("compact")]:
        p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(e["org"]); r.bold = True
        p.add_run(f" — {e['role']}, {e['loc']} ({e['dates']}). "); add_rich(p, ' '.join(e['bullets']))
    heading("Open Source")
    for n, desc, u in OPEN_SOURCE:
        p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1.5)
        for i, (ln, lu) in enumerate(oss_names(n)):
            if i: p.add_run(" · ")
            if lu: add_hyperlink(p, lu, ln)
            else: r = p.add_run(ln); r.bold = True
        p.add_run(" — "); add_rich(p, desc)
    p = d.add_paragraph(); p.add_run(plain(OSS_NOTE)); add_hyperlink(p, OSS_NOTE_URL, OSS_NOTE_LINK_TEXT)
    heading("Publications")
    for n, desc, u in PUBLICATIONS: bullet(desc, bold_prefix=n, url=u)
    heading("Education & Certifications")
    for s, deg, y in EDUCATION: bullet(f"{deg} ({y})", bold_prefix=s)
    for c, o, y in CERTS: bullet(f"{o} ({y})", bold_prefix=c)
    d.core_properties.author = NAME; d.core_properties.title = f"{NAME} — Résumé"
    d.save(path)

# ---------------------------------------------------------------- write
os.makedirs(OUT, exist_ok=True)
open(os.path.join(OUT, "README.md"), "w").write(md())
open(os.path.join(OUT, "adam-steinberger-resume.txt"), "w").write(txt())
open(os.path.join(OUT, "adam-steinberger-resume.html"), "w").write(html_doc())
docx_doc(os.path.join(OUT, "adam-steinberger-resume.docx"))
print("wrote README.md, .txt, .html, .docx →", OUT)
